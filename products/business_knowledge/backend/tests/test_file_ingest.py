import io
import zipfile

import pytest
from posthog.test.base import APIBaseTest
from unittest.mock import patch

from parameterized import parameterized
from rest_framework import status as http_status

from products.business_knowledge.backend.facade.enums import MAX_FILE_SIZE_BYTES
from products.business_knowledge.backend.file_parse import (
    EncryptedPDFError,
    FileParseError,
    FileTooLargeError,
    UnsupportedFileTypeError,
    ZipBombError,
    detect_content_type,
    parse_file,
    sanitize_filename,
)
from products.business_knowledge.backend.models import KnowledgeChunk, KnowledgeDocument, KnowledgeSource

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_minimal_pdf(text: str = "Hello from page 1", encrypted: bool = False) -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    # pypdf doesn't expose a nice text-write API for new pages, so we write a
    # second page with annotation-style text via the low-level API. For tests
    # it's simpler to create a real PDF via reportlab, but we avoid that dep.
    # Instead just test with the blank page — extract_text returns "".
    # We'll test content extraction separately.

    buf = io.BytesIO()
    if encrypted:
        writer.encrypt("secret")
    writer.write(buf)
    return buf.getvalue()


def _make_minimal_docx(text: str = "Hello World") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# detect_content_type
# ---------------------------------------------------------------------------


class TestDetectContentType:
    def test_pdf_from_magic_bytes(self) -> None:
        assert detect_content_type(b"%PDF-1.4 ...", "mystery.bin") == "application/pdf"

    def test_docx_from_zip_structure(self) -> None:
        data = _make_minimal_docx()
        assert detect_content_type(data, "doc.docx") == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    @parameterized.expand(
        [
            ("readme.md", "text/markdown"),
            ("README.MARKDOWN", "text/markdown"),
            ("data.csv", "text/csv"),
            ("notes.txt", "text/plain"),
        ]
    )
    def test_text_formats_from_extension(self, filename: str, expected: str) -> None:
        assert detect_content_type(b"some text content", filename) == expected

    def test_unknown_utf8_falls_back_to_plain_text(self) -> None:
        assert detect_content_type(b"just utf8", "noext") == "text/plain"

    def test_binary_garbage_rejected(self) -> None:
        with pytest.raises(UnsupportedFileTypeError):
            detect_content_type(b"\x89PNG\r\n\x1a\n", "image.png")

    def test_mimetype_spoof_pdf_extension_but_text_content(self) -> None:
        # A .pdf extension with text content should be detected as text/plain
        # from magic bytes (no %PDF- prefix), falling through to extension
        # check. Since .pdf isn't in the text extension list, it falls to
        # UTF-8 check and becomes text/plain.
        result = detect_content_type(b"This is not a PDF", "fake.pdf")
        assert result == "text/plain"

    def test_mimetype_spoof_txt_extension_but_pdf_content(self) -> None:
        pdf_data = _make_minimal_pdf()
        result = detect_content_type(pdf_data, "sneaky.txt")
        assert result == "application/pdf"


# ---------------------------------------------------------------------------
# sanitize_filename
# ---------------------------------------------------------------------------


class TestSanitizeFilename:
    def test_strips_path_components(self) -> None:
        assert sanitize_filename("/etc/passwd") == "passwd"
        assert sanitize_filename("..\\..\\secret.txt") == "secret.txt"

    def test_removes_null_bytes(self) -> None:
        assert sanitize_filename("file\x00name.txt") == "filename.txt"

    def test_caps_length(self) -> None:
        long_name = "a" * 300 + ".pdf"
        assert len(sanitize_filename(long_name)) <= 255

    def test_empty_becomes_unnamed(self) -> None:
        assert sanitize_filename("") == "unnamed"


# ---------------------------------------------------------------------------
# parse_file — happy paths
# ---------------------------------------------------------------------------


class TestParseFileHappyPaths:
    def test_parse_txt(self) -> None:
        result = parse_file(b"Hello, world!\n\nSecond paragraph.", "notes.txt")
        assert result.content_type == "text/plain"
        assert "Hello, world!" in result.content
        assert result.title == "notes"
        assert result.metadata["file_type"] == "txt"

    def test_parse_markdown(self) -> None:
        md = b"# Heading\n\nSome text\n\n## Sub\n\nMore text"
        result = parse_file(md, "readme.md")
        assert result.content_type == "text/markdown"
        assert "# Heading" in result.content
        assert result.metadata["file_type"] == "markdown"

    def test_parse_csv(self) -> None:
        csv_data = b"name,age\nAlice,30\nBob,25"
        result = parse_file(csv_data, "people.csv")
        assert result.content_type == "text/csv"
        assert "name: Alice" in result.content
        assert "age: 30" in result.content
        assert result.metadata["row_count"] == 2

    def test_parse_docx(self) -> None:
        data = _make_minimal_docx("Test paragraph")
        result = parse_file(data, "doc.docx")
        assert result.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert "Test paragraph" in result.content
        assert "# Test Heading" in result.content
        assert result.metadata["file_type"] == "docx"


# ---------------------------------------------------------------------------
# parse_file — security rejections
# ---------------------------------------------------------------------------


class TestParseFileSecurity:
    def test_oversized_file_rejected(self) -> None:
        data = b"x" * (MAX_FILE_SIZE_BYTES + 1)
        with pytest.raises(FileTooLargeError):
            parse_file(data, "huge.txt")

    def test_encrypted_pdf_rejected(self) -> None:
        data = _make_minimal_pdf(encrypted=True)
        with pytest.raises(EncryptedPDFError, match="Encrypted"):
            parse_file(data, "secret.pdf")

    def test_zip_bomb_docx_rejected(self) -> None:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as zf:
            zf.writestr("word/document.xml", "A" * 200)
        data = buf.getvalue()
        with patch("products.business_knowledge.backend.file_parse.MAX_FILE_DECOMPRESSED_BYTES", 50):
            with pytest.raises(ZipBombError):
                parse_file(data, "bomb.docx")

    def test_unsupported_binary_rejected(self) -> None:
        with pytest.raises(UnsupportedFileTypeError):
            parse_file(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100, "image.png")

    def test_empty_txt_rejected(self) -> None:
        with pytest.raises(FileParseError, match="empty"):
            parse_file(b"   \n  ", "empty.txt")

    def test_empty_csv_rejected(self) -> None:
        with pytest.raises(FileParseError, match="empty"):
            parse_file(b"", "empty.csv")

    def test_csv_max_rows_truncated(self) -> None:
        header = "col1,col2\n"
        rows = "".join(f"val{i},data{i}\n" for i in range(15_000))
        data = (header + rows).encode()
        result = parse_file(data, "big.csv")
        assert "Truncated" in result.content


# ---------------------------------------------------------------------------
# API integration
# ---------------------------------------------------------------------------


class TestFileSourceAPIIntegration(APIBaseTest):
    def setUp(self) -> None:
        super().setUp()
        self.url = f"/api/environments/{self.team.id}/business_knowledge/sources/"

    def test_upload_txt_file(self) -> None:
        content = b"Hello world\n\nSecond paragraph about pricing."
        uploaded = io.BytesIO(content)
        uploaded.name = "notes.txt"
        response = self.client.post(
            self.url,
            {"name": "My notes", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_201_CREATED, response.content
        body = response.json()
        assert body["source_type"] == "file"
        assert body["status"] == "ready"
        assert body["original_filename"] == "notes.txt"
        assert body["file_content_type"] == "text/plain"
        assert body["chunk_count"] >= 1

        source = KnowledgeSource.objects.get(id=body["id"])
        assert KnowledgeDocument.objects.filter(source=source, team=self.team).count() == 1
        assert KnowledgeChunk.objects.filter(source=source, team=self.team).count() >= 1

    def test_upload_csv_file(self) -> None:
        content = b"name,role\nAlice,Engineer\nBob,Designer"
        uploaded = io.BytesIO(content)
        uploaded.name = "team.csv"
        response = self.client.post(
            self.url,
            {"name": "Team roster", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_201_CREATED
        body = response.json()
        assert body["file_content_type"] == "text/csv"
        assert body["chunk_count"] >= 1

    def test_upload_markdown_file(self) -> None:
        content = b"# FAQ\n\n## What is PostHog?\n\nPostHog is a product analytics platform."
        uploaded = io.BytesIO(content)
        uploaded.name = "faq.md"
        response = self.client.post(
            self.url,
            {"name": "FAQ", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_201_CREATED
        body = response.json()
        assert body["file_content_type"] == "text/markdown"

    def test_upload_docx_file(self) -> None:
        data = _make_minimal_docx("DOCX test content")
        uploaded = io.BytesIO(data)
        uploaded.name = "document.docx"
        response = self.client.post(
            self.url,
            {"name": "Word doc", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_201_CREATED
        body = response.json()
        assert body["file_content_type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert body["chunk_count"] >= 1

    def test_upload_rejects_unsupported_type(self) -> None:
        uploaded = io.BytesIO(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
        uploaded.name = "image.png"
        response = self.client.post(
            self.url,
            {"name": "Image", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_400_BAD_REQUEST

    def test_upload_rejects_oversized_file(self) -> None:
        # The serializer checks UploadedFile.size which Django sets from the
        # actual content length. We can't send 50 MB in a unit test, but the
        # parse_file size guard is tested separately in TestParseFileSecurity.
        # Here we verify the parse_file guard fires at the API layer by
        # monkey-patching the cap to a tiny value.
        with patch("products.business_knowledge.backend.file_parse.MAX_FILE_SIZE_BYTES", 10):
            uploaded = io.BytesIO(b"x" * 100)
            uploaded.name = "big.txt"
            response = self.client.post(
                self.url,
                {"name": "Too big", "file": uploaded, "source_type": "file"},
                format="multipart",
            )
        assert response.status_code == http_status.HTTP_400_BAD_REQUEST

    def test_upload_missing_file_field(self) -> None:
        response = self.client.post(
            self.url,
            {"name": "No file", "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_400_BAD_REQUEST

    def test_upload_missing_name(self) -> None:
        uploaded = io.BytesIO(b"content")
        uploaded.name = "data.txt"
        response = self.client.post(
            self.url,
            {"file": uploaded, "source_type": "file"},
            format="multipart",
        )
        assert response.status_code == http_status.HTTP_400_BAD_REQUEST

    def test_file_source_appears_in_list(self) -> None:
        uploaded = io.BytesIO(b"test content")
        uploaded.name = "test.txt"
        self.client.post(
            self.url,
            {"name": "Listed file", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        response = self.client.get(self.url)
        names = [r["name"] for r in response.json()["results"]]
        assert "Listed file" in names

    def test_file_source_cross_team_isolation(self) -> None:
        from posthog.models.team import Team

        uploaded = io.BytesIO(b"secret data")
        uploaded.name = "secret.txt"
        self.client.post(
            self.url,
            {"name": "Secret", "file": uploaded, "source_type": "file"},
            format="multipart",
        )
        source = KnowledgeSource.objects.filter(team=self.team).first()
        assert source is not None

        other_team = Team.objects.create_with_data(
            organization=self.organization, initiating_user=self.user, name="Other"
        )
        other_url = f"/api/environments/{other_team.id}/business_knowledge/sources/"
        response = self.client.get(f"{other_url}{source.id}/")
        assert response.status_code == http_status.HTTP_404_NOT_FOUND
